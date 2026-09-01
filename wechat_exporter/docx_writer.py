"""Word 导出模块：把结构化 Article 渲染为排版清晰的 .docx。

排版策略：
  - 标题 → Word Heading 1
  - 元信息（发布时间 / 原文链接 / 摘要）→ 灰色小字信息区
  - 正文 → 宋体小四，1.5 倍行距，首行缩进 2 字符
  - 图片 → 居中，按页宽自适应缩放
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .config import settings
from .parser import Article, Block

logger = logging.getLogger("dwechatword.docx_writer")

# A4 正文可用宽度（页边距 2.5cm 左右各）
MAX_IMAGE_WIDTH_CM = 16.0


def _set_font(run, name_ascii: str = "Times New Roman",
              name_east: str = "宋体", size: float | None = None) -> None:
    run.font.name = name_ascii
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name_east)
    if size:
        run.font.size = Pt(size)


def _sanitize_filename(name: str, max_len: int = 60) -> str:
    name = re.sub(r'[\\/:*?"<>|\r\n\t]', "_", name).strip(" .")
    return name[:max_len] or "untitled"


def build_docx(article: Article, out_dir: Path | None = None) -> Path:
    out_dir = out_dir or settings.output_dir
    out_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 页面默认样式
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)  # 小四

    # 标题
    h = doc.add_heading(article.title, level=1)
    for run in h.runs:
        _set_font(run, "微软雅黑", "微软雅黑")

    # 元信息区
    meta_lines = [f"发布时间：{article.publish_time or '未知'}"]
    if article.digest:
        meta_lines.append(f"摘要：{article.digest}")
    if article.url:
        meta_lines.append(f"原文链接：{article.url}")
    for line in meta_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()  # 空行分隔

    # 正文块
    for block in article.blocks:
        _render_block(doc, block)

    # 文件命名：yyyy-mm-dd-标题-版本号.docx（日期取发布日期）
    date_part = article.publish_time[:10] if len(article.publish_time) >= 10 else "unknown-date"
    filename = (f"{date_part}-{_sanitize_filename(article.title)}"
                f"-{settings.doc_version}.docx")
    out_path = out_dir / filename
    doc.save(out_path)
    logger.info("已导出：%s", out_path.name)
    return out_path


def _render_block(doc: Document, block: Block) -> None:
    if block.type == "image" and block.image_path:
        _render_image(doc, block.image_path)
        return

    if not block.runs:
        return

    if block.type == "heading":
        h = doc.add_heading("".join(r.text for r in block.runs),
                            level=min(block.level + 1, 4))
        for run in h.runs:
            _set_font(run, "微软雅黑", "微软雅黑")
        return

    if block.type == "quote":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        for r in block.runs:
            run = p.add_run(r.text)
            _set_font(run, size=11)
            run.italic = True
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        return

    if block.type == "list_item":
        p = doc.add_paragraph(style="List Bullet")
        _add_runs(p, block.runs)
        return

    # 普通段落
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(24)  # 首行缩进 2 字符
    _add_runs(p, block.runs)


def _add_runs(paragraph, runs) -> None:
    for r in runs:
        run = paragraph.add_run(r.text)
        _set_font(run, size=12)
        run.bold = r.bold
        run.italic = r.italic
        run.underline = r.underline


def _render_image(doc: Document, image_path: Path) -> None:
    try:
        from PIL import Image
        with Image.open(image_path) as im:
            width_px, height_px = im.size
        # 图片实际尺寸（微信 CDN 常带 wx_width 信息，简单按 640px 宽度基准换算）
        width_cm = min(width_px / 640 * 16.0, MAX_IMAGE_WIDTH_CM)
        width_cm = max(width_cm, 4.0)
    except Exception:  # noqa: BLE001 读取失败按默认宽度
        width_cm = MAX_IMAGE_WIDTH_CM
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    except Exception as e:  # noqa: BLE001
        logger.warning("图片插入 Word 失败 %s: %s", image_path, e)
